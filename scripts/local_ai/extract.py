#!/usr/bin/env python3
"""Local document extraction. Everything here runs on this machine.

The cheapest method that works is always tried first, because handing a raw
scanned PDF to a vision model for every page is slow and less accurate than
reading a text layer that is already there.

    native text  ->  tables  ->  OCR  ->  page images for a vision model

    ./vendor/hermes-venv/bin/python scripts/local_ai/extract.py <file>
    ./vendor/hermes-venv/bin/python scripts/local_ai/extract.py <file> --json
    ./vendor/hermes-venv/bin/python scripts/local_ai/extract.py <file> --classify-only

Supported: PDF (native text and scanned), DOCX, XLSX, and images.

Nothing in this module makes a network call. That is deliberate and is what
lets Local Privacy Mode make an honest promise about extraction.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "lib"))
from miniyaml import load_file  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent.parent
CONFIG = ROOT / "config" / "local-ai.yaml"
CONFIG_EXAMPLE = ROOT / "config" / "local-ai.example.yaml"
# Rendered page images for the vision model. Gitignored, stays local.
WORKING = ROOT / "local_data" / "working" / "pages"

# Signals used to classify a document from its own text.
#
# Markers are WEIGHTED because generic words are actively misleading. A Closing
# Disclosure legitimately contains "seller", "buyer", "earnest money", and
# "closing date" — the same words a purchase contract has — so unweighted
# counting produces a tie that resolves by dictionary order. Weighting the
# near-unique title phrases well above generic terms fixes that.
#
#   3 = near-unique to this document type (usually the form's own title)
#   2 = strong signal
#   1 = weak / shared with other document types
DOC_SIGNATURES = [
    ("closing_disclosure", {
        "closing disclosure": 3, "summaries of transactions": 3,
        "calculating cash to close": 3, "disbursement date": 2,
        "total closing costs": 2, "lender credits": 1, "cash to close": 1,
        "loan costs": 1,
    }),
    ("loan_estimate", {
        "loan estimate": 3, "estimated cash to close": 3,
        "closing cost details": 2, "projected payments": 2,
        "services you can shop for": 2, "loan terms": 1, "cash to close": 1,
    }),
    ("paystub", {
        "earnings statement": 3, "pay stub": 3, "paystub": 3,
        "year to date": 2, "gross pay": 2, "net pay": 2, "pay period": 2,
        "ytd": 1, "deductions": 1,
    }),
    ("w2", {
        "wage and tax statement": 3, "form w-2": 3, "w-2": 2,
        "social security wages": 2, "medicare wages": 2,
        "federal income tax withheld": 1,
    }),
    ("tax_return", {
        "u.s. individual income tax return": 3, "form 1040": 3,
        "adjusted gross income": 2, "schedule c": 2, "schedule e": 2,
        "schedule se": 2, "filing status": 1, "taxable income": 1,
    }),
    ("bank_statement", {
        "account summary": 3, "beginning balance": 3, "ending balance": 3,
        "statement period": 2, "deposits and additions": 2,
        "average daily balance": 2, "withdrawals": 1,
    }),
    ("purchase_contract", {
        "purchase and sale": 3, "residential purchase agreement": 3,
        "purchase agreement": 3, "inspection period": 2,
        "financing contingency": 2, "appraisal contingency": 2,
        "purchase price": 2, "earnest money": 1, "closing date": 1,
        "seller": 1, "buyer": 1,
    }),
    ("mortgage_statement", {
        "mortgage statement": 3, "payment due date": 2,
        "outstanding principal balance": 2, "principal balance": 2,
        "escrow balance": 2, "amount due": 1,
    }),
    ("appraisal", {
        "uniform residential appraisal": 3, "appraised value": 3,
        "comparable sale": 2, "subject property": 2, "gross living area": 2,
    }),
    ("insurance", {
        "declarations page": 3, "dwelling coverage": 3, "policy number": 2,
        "premium": 1, "coverage": 1,
    }),
    ("hoa", {
        "homeowners association": 3, "covenants": 2, "hoa": 2,
        "assessment": 1,
    }),
    ("1099", {
        "form 1099": 3, "1099-misc": 3, "1099-nec": 3,
        "nonemployee compensation": 2,
    }),
]


@dataclass
class PageResult:
    page: int
    method: str                      # native | ocr | image-only
    char_count: int
    text: str = ""
    needs_vision: bool = False
    # Set when a page has no readable text layer and OCR could not help. The
    # rendered image is written to local_data/working/ so the local vision
    # model can read it. It never leaves the machine.
    image_path: str = ""


@dataclass
class ExtractionResult:
    source_file: str
    file_type: str
    processed_locally: bool = True
    document_type: str = "unknown"
    classification_confidence: str = "low"
    page_count: int = 0
    pages: list = field(default_factory=list)
    tables: list = field(default_factory=list)
    warnings: list = field(default_factory=list)
    methods_used: list = field(default_factory=list)
    full_text: str = ""


def load_config() -> dict:
    return load_file(CONFIG if CONFIG.exists() else CONFIG_EXAMPLE) or {}


def classify(text: str) -> tuple[str, str]:
    """Guess the document type from its text. Returns (type, confidence).

    Scores each type by summing the weights of the markers it matches, then
    requires a clear margin over the runner-up before claiming high confidence.
    """
    low = (text or "").lower()
    if not low.strip():
        return "unknown", "none"

    scores = {}
    for doc_type, markers in DOC_SIGNATURES:
        score = sum(weight for marker, weight in markers.items() if marker in low)
        if score:
            scores[doc_type] = score
    if not scores:
        return "unknown", "none"

    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    best, best_score = ranked[0]
    runner_up = ranked[1][1] if len(ranked) > 1 else 0
    margin = best_score - runner_up

    # A high-weight match is a form title; that plus a clear margin is decisive.
    if best_score >= 5 and margin >= 2:
        return best, "high"
    if best_score >= 3 and margin >= 1:
        return best, "moderate"
    if best_score >= 3:
        # Tied or nearly tied — say so rather than guessing confidently.
        return best, "low"
    return best, "low"


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def render_page(page, source: Path, number: int, dpi: int) -> str:
    """Render one PDF page to a local PNG for the vision model.

    Written under local_data/working/, which is gitignored. The image never
    leaves this machine — it is read back and passed to the local model only.
    """
    try:
        WORKING.mkdir(parents=True, exist_ok=True)
        out = WORKING / f"{source.stem}-p{number}.png"
        page.get_pixmap(dpi=dpi).save(out)
        return str(out)
    except Exception:  # noqa: BLE001 - a failed render is reported, not fatal
        return ""


def extract_pdf(path: Path, config: dict) -> ExtractionResult:
    import pymupdf

    docs_cfg = config.get("documents") or {}
    ocr_cfg = docs_cfg.get("ocr") or {}
    min_native = int(docs_cfg.get("min_chars_for_native") or ocr_cfg.get("min_chars_for_native") or 100)
    max_pages = int(docs_cfg.get("max_pages_per_document") or 50)
    dpi = int(docs_cfg.get("render_dpi") or 200)
    ocr_enabled = bool(ocr_cfg.get("enabled", True))

    result = ExtractionResult(source_file=str(path.name), file_type="pdf")
    doc = pymupdf.open(path)
    result.page_count = doc.page_count

    if doc.page_count > max_pages:
        result.warnings.append(
            f"Document has {doc.page_count} pages; only the first {max_pages} were processed "
            f"(max_pages_per_document in config/local-ai.yaml)."
        )

    ocr_available = bool(shutil.which("tesseract"))
    if ocr_enabled and not ocr_available:
        result.warnings.append(
            "OCR is enabled but tesseract is not installed. Scanned pages could not be read. "
            "Install it with: brew install tesseract"
        )

    for index in range(min(doc.page_count, max_pages)):
        page = doc[index]
        text = page.get_text().strip()

        if len(text) >= min_native:
            result.pages.append(PageResult(index + 1, "native", len(text), text))
            continue

        # Not enough text: this page is scanned or image-only.
        if ocr_enabled and ocr_available:
            try:
                import pytesseract
                from PIL import Image
                import io

                pix = page.get_pixmap(dpi=dpi)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(image).strip()
                result.pages.append(PageResult(index + 1, "ocr", len(ocr_text), ocr_text))
                if len(ocr_text) < min_native:
                    result.warnings.append(
                        f"Page {index + 1}: OCR produced very little text ({len(ocr_text)} chars). "
                        f"The scan may be poor quality — verify this page by hand."
                    )
            except Exception as exc:  # noqa: BLE001 - report, never fall back to cloud
                image = render_page(page, path, index + 1, dpi)
                result.pages.append(PageResult(index + 1, "image-only", 0, "",
                                               needs_vision=True, image_path=image))
                result.warnings.append(f"Page {index + 1}: local OCR failed ({exc}). "
                                       f"Handing this page to the local vision model.")
        else:
            image = render_page(page, path, index + 1, dpi)
            result.pages.append(PageResult(index + 1, "image-only", 0, "",
                                           needs_vision=True, image_path=image))

    # Tables, from pages that had a real text layer.
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages[:max_pages]):
                for table in page.extract_tables() or []:
                    cleaned = [[(c or "").strip() for c in row] for row in table if any(row)]
                    if len(cleaned) > 1:
                        result.tables.append({"page": index + 1, "rows": cleaned})
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"Table extraction did not run: {exc}")

    doc.close()
    _finalize(result)
    return result


# --------------------------------------------------------------------------
# Other formats
# --------------------------------------------------------------------------

def extract_docx(path: Path, config: dict) -> ExtractionResult:
    import docx

    result = ExtractionResult(source_file=path.name, file_type="docx")
    document = docx.Document(path)
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    result.pages.append(PageResult(1, "native", len(text), text))
    result.page_count = 1
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) > 1:
            result.tables.append({"page": 1, "rows": rows})
    _finalize(result)
    return result


def extract_xlsx(path: Path, config: dict) -> ExtractionResult:
    import openpyxl

    result = ExtractionResult(source_file=path.name, file_type="xlsx")
    book = openpyxl.load_workbook(path, data_only=True)
    chunks = []
    for sheet in book.worksheets:
        rows = [[("" if c is None else str(c)) for c in row]
                for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if any(v.strip() for v in r)]
        if rows:
            result.tables.append({"sheet": sheet.title, "rows": rows})
            chunks.append(f"[{sheet.title}]\n" + "\n".join("\t".join(r) for r in rows))
    text = "\n\n".join(chunks)
    result.pages.append(PageResult(1, "native", len(text), text))
    result.page_count = 1
    _finalize(result)
    return result


def extract_image(path: Path, config: dict) -> ExtractionResult:
    result = ExtractionResult(source_file=path.name, file_type="image")
    result.page_count = 1
    ocr_cfg = (config.get("documents") or {}).get("ocr") or {}
    if ocr_cfg.get("enabled", True) and shutil.which("tesseract"):
        try:
            import pytesseract
            from PIL import Image
            text = pytesseract.image_to_string(Image.open(path)).strip()
            result.pages.append(PageResult(1, "ocr", len(text), text))
        except Exception as exc:  # noqa: BLE001
            result.pages.append(PageResult(1, "image-only", 0, "", needs_vision=True,
                                           image_path=str(path)))
            result.warnings.append(f"Local OCR failed ({exc}). Handing this to the local vision model.")
    else:
        result.pages.append(PageResult(1, "image-only", 0, "", needs_vision=True,
                                       image_path=str(path)))
        result.warnings.append("OCR unavailable — handing this image to the local vision model.")
    _finalize(result)
    return result


def _finalize(result: ExtractionResult) -> None:
    result.full_text = "\n\n".join(p.text for p in result.pages if p.text)
    result.methods_used = sorted({p.method for p in result.pages})
    result.document_type, result.classification_confidence = classify(result.full_text)
    if any(p.needs_vision for p in result.pages):
        pending = [p.page for p in result.pages if p.needs_vision]
        result.warnings.append(
            f"Page(s) {pending} have no readable text layer and will be read by the "
            f"local vision model. They were NOT sent anywhere off this machine."
        )
    total = sum(p.char_count for p in result.pages)
    if total < 50:
        result.warnings.append(
            "Almost no text was extracted. Confirm the file is a real document and "
            "not an empty or corrupt PDF."
        )


EXTRACTORS = {
    ".pdf": extract_pdf, ".docx": extract_docx, ".xlsx": extract_xlsx, ".xlsm": extract_xlsx,
    ".png": extract_image, ".jpg": extract_image, ".jpeg": extract_image,
    ".tif": extract_image, ".tiff": extract_image, ".webp": extract_image,
}


def extract(path: Path) -> ExtractionResult:
    config = load_config()
    suffix = path.suffix.lower()
    handler = EXTRACTORS.get(suffix)
    if not handler:
        raise SystemExit(
            f"Unsupported file type {suffix!r}. Supported: {', '.join(sorted(EXTRACTORS))}"
        )
    if not path.exists():
        raise SystemExit(f"No such file: {path}")
    return handler(path, config)


def to_dict(result: ExtractionResult) -> dict:
    data = asdict(result)
    data["pages"] = [asdict(p) if not isinstance(p, dict) else p for p in result.pages]
    return data


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract a document locally.")
    parser.add_argument("file")
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--classify-only", action="store_true")
    parser.add_argument("--max-chars", type=int, default=1500)
    args = parser.parse_args()

    result = extract(Path(args.file))

    if args.json:
        print(json.dumps(to_dict(result), indent=2))
        return 0

    print(f"\n  Local extraction — {result.source_file}\n  " + "-" * 50)
    print(f"  File type      : {result.file_type}")
    print(f"  Pages          : {result.page_count}")
    print(f"  Methods used   : {', '.join(result.methods_used) or 'none'}")
    print(f"  Document type  : {result.document_type} (confidence: {result.classification_confidence})")
    print(f"  Tables found   : {len(result.tables)}")
    print(f"  Processed      : entirely on this machine")

    if result.warnings:
        print("\n  Warnings")
        for warning in result.warnings:
            print(f"    - {warning}")

    if not args.classify_only and result.full_text:
        print("\n  Extracted text (first "
              f"{args.max_chars} chars)\n  " + "-" * 50)
        for line in result.full_text[:args.max_chars].splitlines():
            print(f"  {line}")
    print()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
